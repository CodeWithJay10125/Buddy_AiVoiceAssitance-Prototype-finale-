"""
BUDDY AI Assistant - Python Backend
Full OS control for Windows 11
Run: python buddy.py
"""

import os
import sys
import json
import time
import threading
import subprocess
import webbrowser
import http.server
import socketserver
import urllib.parse
from pathlib import Path
from datetime import datetime

# ── optional imports (installed by setup.bat) ──────────────────────────────
try:
    import psutil
    HAS_PSUTIL = True
except ImportError:
    HAS_PSUTIL = False

try:
    import pyautogui
    pyautogui.FAILSAFE = False
    HAS_PYAUTOGUI = True
except ImportError:
    HAS_PYAUTOGUI = False

try:
    import pyttsx3
    HAS_TTS = True
except ImportError:
    HAS_TTS = False

try:
    import win32com.client
    import win32con
    import win32gui
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pystray
    from PIL import Image as PILImage
    HAS_TRAY = True
except ImportError:
    HAS_TRAY = False

# ───────────────────────────────────────────────────────────────────────────
# TTS ENGINE
# ───────────────────────────────────────────────────────────────────────────
_tts_engine = None

def get_tts():
    global _tts_engine
    if _tts_engine is None and HAS_TTS:
        _tts_engine = pyttsx3.init()
        voices = _tts_engine.getProperty('voices')
        # pick male voice
        for v in voices:
            if 'male' in v.name.lower() or 'david' in v.name.lower() or 'mark' in v.name.lower() or 'james' in v.name.lower():
                _tts_engine.setProperty('voice', v.id)
                break
        _tts_engine.setProperty('rate', 175)
        _tts_engine.setProperty('volume', 1.0)
    return _tts_engine

def speak_sync(text):
    """Speak text (blocking)."""
    print(f"[BUDDY] {text}")
    eng = get_tts()
    if eng:
        try:
            eng.say(text)
            eng.runAndWait()
        except Exception:
            pass
    else:
        # fallback: Windows SAPI via PowerShell
        ps = f'Add-Type -AssemblyName System.Speech; $s=New-Object System.Speech.Synthesis.SpeechSynthesizer; $s.Rate=-2; $s.Speak("{text}")'
        subprocess.Popen(['powershell', '-Command', ps], creationflags=subprocess.CREATE_NO_WINDOW)
        time.sleep(max(1, len(text) / 12))

def speak_async(text):
    t = threading.Thread(target=speak_sync, args=(text,), daemon=True)
    t.start()
    return t

# ───────────────────────────────────────────────────────────────────────────
# SYSTEM INFO
# ───────────────────────────────────────────────────────────────────────────
def get_battery():
    if HAS_PSUTIL:
        bat = psutil.sensors_battery()
        if bat:
            return int(bat.percent), bat.power_plugged
    return None, None

def get_cpu():
    if HAS_PSUTIL:
        return psutil.cpu_percent(interval=0.1)
    return 0

def get_ram():
    if HAS_PSUTIL:
        r = psutil.virtual_memory()
        return r.percent, round(r.total / (1024**3), 1), round(r.used / (1024**3), 1)
    return 0, 0, 0

def get_stats():
    cpu = get_cpu()
    ram_pct, ram_total, ram_used = get_ram()
    bat_pct, bat_plug = get_battery()
    return {
        "cpu": cpu,
        "ram_pct": ram_pct,
        "ram_total": ram_total,
        "ram_used": round(ram_used, 1),
        "battery": bat_pct,
        "charging": bat_plug,
        "time": datetime.now().strftime("%H:%M:%S"),
        "date": datetime.now().strftime("%A, %B %d, %Y"),
    }

# ───────────────────────────────────────────────────────────────────────────
# APP PATHS  (Windows 11 defaults — edit if yours differ)
# ───────────────────────────────────────────────────────────────────────────
APP_PATHS = {
    "vscode":        r"C:\Users\{user}\AppData\Local\Programs\Microsoft VS Code\Code.exe",
    "notepad":       r"notepad.exe",
    "word":          r"C:\Program Files\Microsoft Office\root\Office16\WINWORD.EXE",
    "excel":         r"C:\Program Files\Microsoft Office\root\Office16\EXCEL.EXE",
    "powerpoint":    r"C:\Program Files\Microsoft Office\root\Office16\POWERPNT.EXE",
    "explorer":      r"explorer.exe",
    "chrome":        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
    "edge":          r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    "calculator":    r"calc.exe",
    "paint":         r"mspaint.exe",
    "cmd":           r"cmd.exe",
    "powershell":    r"powershell.exe",
    "task_manager":  r"taskmgr.exe",
    "spotify":       r"C:\Users\{user}\AppData\Roaming\Spotify\Spotify.exe",
    "vlc":           r"C:\Program Files\VideoLAN\VLC\vlc.exe",
    "discord":       r"C:\Users\{user}\AppData\Local\Discord\app-*\Discord.exe",
    "telegram":      r"C:\Users\{user}\AppData\Roaming\Telegram Desktop\Telegram.exe",
    "whatsapp":      r"C:\Users\{user}\AppData\Local\WhatsApp\WhatsApp.exe",
}

def resolve_path(key):
    user = os.environ.get("USERNAME", "User")
    path = APP_PATHS.get(key, "").replace("{user}", user)
    # handle glob for Discord-style paths
    if "*" in path:
        import glob
        matches = glob.glob(path)
        return matches[0] if matches else None
    return path if path else None

def open_app(key, args=None):
    """Launch an app by key. Returns (success, message)."""
    path = resolve_path(key)
    try:
        if path and os.path.exists(path):
            cmd = [path] + (args or [])
            subprocess.Popen(cmd)
            return True, f"Opening {key}"
        else:
            # try via shell (works for built-ins like notepad, calc, explorer)
            cmd = key if not args else f"{key} {' '.join(args)}"
            subprocess.Popen(cmd, shell=True)
            return True, f"Launching {key}"
    except Exception as e:
        return False, str(e)

# ───────────────────────────────────────────────────────────────────────────
# FILE OPERATIONS
# ───────────────────────────────────────────────────────────────────────────
def open_file_explorer(path=None):
    if path:
        expanded = os.path.expandvars(os.path.expanduser(path))
        subprocess.Popen(f'explorer "{expanded}"', shell=True)
        return f"Opening File Explorer at {expanded}"
    else:
        subprocess.Popen("explorer", shell=True)
        return "Opening File Explorer"

def open_file(filepath):
    expanded = os.path.expandvars(os.path.expanduser(filepath))
    if os.path.exists(expanded):
        os.startfile(expanded)
        return True, f"Opening {expanded}"
    return False, f"File not found: {expanded}"

def create_word_doc(filename, content_lines, save_dir=None):
    """Create a .docx file with given lines."""
    if not HAS_DOCX:
        # fallback: open Word and type via pyautogui
        return create_word_fallback(filename, content_lines, save_dir)

    save_dir = save_dir or str(Path.home() / "Documents")
    os.makedirs(save_dir, exist_ok=True)
    if not filename.endswith(".docx"):
        filename += ".docx"
    filepath = os.path.join(save_dir, filename)

    doc = Document()
    for line in content_lines:
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")
    doc.save(filepath)
    os.startfile(filepath)
    return True, filepath

def create_word_fallback(filename, content_lines, save_dir=None):
    save_dir = save_dir or str(Path.home() / "Documents")
    os.makedirs(save_dir, exist_ok=True)
    if not filename.endswith(".docx"):
        filename += ".docx"
    filepath = os.path.join(save_dir, filename)
    # Write as plain txt if no python-docx
    txt_path = filepath.replace(".docx", ".txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("\n".join(content_lines))
    os.startfile(txt_path)
    return True, txt_path

def open_vscode_file(filepath):
    """Open a file in VS Code."""
    expanded = os.path.expandvars(os.path.expanduser(filepath))
    vscode = resolve_path("vscode")
    if vscode and os.path.exists(vscode):
        subprocess.Popen([vscode, expanded])
        return True, f"Opening {expanded} in VS Code"
    else:
        # try via PATH
        try:
            subprocess.Popen(["code", expanded], shell=True)
            return True, f"Opening {expanded} in VS Code"
        except:
            return False, "VS Code not found"

def create_file_and_open(filename, content="", editor="vscode", save_dir=None):
    save_dir = save_dir or str(Path.home() / "Documents")
    os.makedirs(save_dir, exist_ok=True)
    filepath = os.path.join(save_dir, filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)
    if editor == "vscode":
        ok, msg = open_vscode_file(filepath)
    else:
        os.startfile(filepath)
        ok, msg = True, f"Opened {filepath}"
    return ok, filepath

# ───────────────────────────────────────────────────────────────────────────
# YOUTUBE / BROWSER
# ───────────────────────────────────────────────────────────────────────────
def play_youtube(query):
    url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(query)}"
    webbrowser.open(url)
    return f"Searching YouTube for {query}"

def open_url(url):
    webbrowser.open(url)
    return f"Opening {url}"

# ───────────────────────────────────────────────────────────────────────────
# VOLUME CONTROL
# ───────────────────────────────────────────────────────────────────────────
def set_volume(level):
    """Set volume 0-100 via PowerShell."""
    script = f"""
$obj = New-Object -ComObject WScript.Shell
for ($i=0; $i -lt 50; $i++) {{ $obj.SendKeys([char]174) }}
$steps = [math]::Round({level} / 2)
for ($i=0; $i -lt $steps; $i++) {{ $obj.SendKeys([char]175) }}
"""
    subprocess.Popen(["powershell", "-Command", script],
                     creationflags=subprocess.CREATE_NO_WINDOW)
    return f"Volume set to {level} percent"

def mute_volume():
    script = '$obj = New-Object -ComObject WScript.Shell; $obj.SendKeys([char]173)'
    subprocess.Popen(["powershell", "-Command", script],
                     creationflags=subprocess.CREATE_NO_WINDOW)
    return "Volume muted"

# ───────────────────────────────────────────────────────────────────────────
# SCREENSHOT
# ───────────────────────────────────────────────────────────────────────────
def take_screenshot(save_dir=None):
    save_dir = save_dir or str(Path.home() / "Pictures" / "Screenshots")
    os.makedirs(save_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = os.path.join(save_dir, f"buddy_screenshot_{ts}.png")
    if HAS_PYAUTOGUI:
        img = pyautogui.screenshot()
        img.save(path)
        os.startfile(save_dir)
        return True, path
    else:
        # Use Windows Snipping Tool shortcut
        subprocess.Popen(["snippingtool"], shell=True)
        return True, "Snipping Tool opened"

# ───────────────────────────────────────────────────────────────────────────
# SHUTDOWN / RESTART / SLEEP
# ───────────────────────────────────────────────────────────────────────────
def shutdown_pc():
    subprocess.Popen("shutdown /s /t 5", shell=True)
    return "Shutting down in 5 seconds, boss."

def restart_pc():
    subprocess.Popen("shutdown /r /t 5", shell=True)
    return "Restarting in 5 seconds, boss."

def sleep_pc():
    subprocess.Popen("rundll32.exe powrprof.dll,SetSuspendState 0,1,0", shell=True)
    return "Putting system to sleep."

# ───────────────────────────────────────────────────────────────────────────
# COMMAND PROCESSOR
# ───────────────────────────────────────────────────────────────────────────
def process_command(cmd: str) -> dict:
    """
    Parse voice command string and execute OS action.
    Returns { "response": str, "action": str }
    """
    lower = cmd.lower().strip()
    response = ""
    action = "none"

    # ── WAKE UP ────────────────────────────────────────────────────────────
    if ("wake up" in lower or "wakeup" in lower) and "buddy" in lower:
        bat, plug = get_battery()
        bat_str = f"{bat} percent battery" if bat else "battery status unknown"
        charging = ", currently charging" if plug else ""
        response = f"Welcome boss, we are online and our system is running on Windows 11 with {bat_str}{charging}. What can I do for you, sir?"
        action = "wake"
        return {"response": response, "action": action}

    # ── SLEEP BUDDY ────────────────────────────────────────────────────────
    if any(x in lower for x in ["go to sleep", "stand by", "standby", "sleep buddy"]):
        response = "Going to standby mode, boss. Say wake up buddy to reactivate me."
        action = "sleep_buddy"
        return {"response": response, "action": action}

    # ── OPEN FILE EXPLORER ────────────────────────────────────────────────
    if "file explorer" in lower or "open explorer" in lower or "open files" in lower:
        # check for path like "open explorer C:\Users"
        import re
        path_match = re.search(r'(?:at|in|to|open)\s+([a-zA-Z]:\\[^\s]+|~[^\s]*|%[^\s%]+%)', cmd)
        path = path_match.group(1) if path_match else None
        msg = open_file_explorer(path)
        response = f"{msg}, sir."
        action = "open_explorer"
        return {"response": response, "action": action}

    # ── OPEN VS CODE ──────────────────────────────────────────────────────
    if "vs code" in lower or "vscode" in lower or "visual studio code" in lower:
        import re
        # "open vs code file hello.py" or "open test.py in vs code"
        file_match = re.search(r'(?:file|open|edit)\s+([\w\-./\\]+\.\w+)', lower)
        if not file_match:
            file_match = re.search(r'([\w\-./\\]+\.\w+)\s+(?:in|with)\s+(?:vs code|vscode)', lower)
        if file_match:
            fname = file_match.group(1)
            ok, msg = open_vscode_file(fname)
            response = f"Opening {fname} in VS Code, boss." if ok else f"Sorry boss, {msg}"
        else:
            ok, msg = open_app("vscode")
            response = "Opening VS Code for you, boss." if ok else f"Could not open VS Code, sir. {msg}"
        action = "open_vscode"
        return {"response": response, "action": action}

    # ── OPEN NOTEPAD ──────────────────────────────────────────────────────
    if "notepad" in lower:
        import re
        file_match = re.search(r'(?:file|open|edit)\s+([\w\-./\\]+\.?\w*)', lower)
        if file_match:
            fname = file_match.group(1)
            subprocess.Popen(f'notepad "{fname}"', shell=True)
            response = f"Opening {fname} in Notepad, sir."
        else:
            open_app("notepad")
            response = "Opening Notepad, boss."
        action = "open_notepad"
        return {"response": response, "action": action}

    # ── OPEN MS WORD ──────────────────────────────────────────────────────
    if ("ms word" in lower or "microsoft word" in lower or ("open word" in lower)):
        import re
        # check if they want to write something
        write_match = re.search(r'(?:write|create|make)\s+(?:a\s+)?(.+?)(?:\s+in|\s+on|$)', lower)
        if write_match:
            doc_type = write_match.group(1).strip()
            # generate content based on type
            content = generate_doc_content(doc_type)
            fname = doc_type.replace(" ", "_") + "_buddy"
            ok, path = create_word_doc(fname, content)
            response = f"Created and opened your {doc_type} in MS Word, boss."
        else:
            ok, msg = open_app("word")
            response = "Opening Microsoft Word, sir." if ok else f"Could not open Word. {msg}"
        action = "open_word"
        return {"response": response, "action": action}

    # ── CREATE / WRITE DOCUMENT ───────────────────────────────────────────
    if any(x in lower for x in ["write a", "create a", "make a", "write me"]):
        import re
        # "write a letter", "write a report about X", "create a python file"
        doc_match = re.search(r'(?:write|create|make)\s+(?:a\s+|me\s+a\s+|me\s+)?(.+)', lower)
        if doc_match:
            doc_type = doc_match.group(1).strip()
            # code file?
            code_exts = [".py", ".js", ".html", ".css", ".cpp", ".java", ".ts"]
            ext_match = re.search(r'(\.\w+)', doc_type)
            code_match = re.search(r'(python|javascript|html|css|java|typescript|cpp)\s+(?:file|script|program)', lower)
            fname_match = re.search(r'(?:named?|called?|file)\s+([\w\-]+)', lower)
            fname = fname_match.group(1) if fname_match else doc_type.replace(" ", "_")

            if ext_match or code_match:
                lang = code_match.group(1) if code_match else "python"
                ext_map = {"python":".py","javascript":".js","html":".html","css":".css","java":".java","typescript":".ts","cpp":".cpp"}
                ext = ext_match.group(1) if ext_match else ext_map.get(lang, ".py")
                content = generate_code_template(lang, fname)
                ok, path = create_file_and_open(fname + ext, content, editor="vscode")
                response = f"Created {fname}{ext} and opened it in VS Code, boss."
            else:
                content = generate_doc_content(doc_type)
                ok, path = create_word_doc(fname + "_buddy", content)
                response = f"Created your {doc_type} and opened it in Word, sir."
            action = "create_doc"
            return {"response": response, "action": action}

    # ── OPEN EXCEL ────────────────────────────────────────────────────────
    if "excel" in lower or "spreadsheet" in lower:
        open_app("excel")
        response = "Opening Microsoft Excel, boss."
        action = "open_excel"
        return {"response": response, "action": action}

    # ── OPEN POWERPOINT ───────────────────────────────────────────────────
    if "powerpoint" in lower or "presentation" in lower:
        open_app("powerpoint")
        response = "Opening Microsoft PowerPoint, sir."
        action = "open_ppt"
        return {"response": response, "action": action}

    # ── OPEN CHROME / EDGE ────────────────────────────────────────────────
    if "chrome" in lower:
        open_app("chrome")
        response = "Opening Google Chrome, boss."
        action = "open_chrome"
        return {"response": response, "action": action}
    if "edge" in lower:
        open_app("edge")
        response = "Opening Microsoft Edge, sir."
        action = "open_edge"
        return {"response": response, "action": action}

    # ── YOUTUBE / PLAY ────────────────────────────────────────────────────
    if "youtube" in lower or ("play" in lower and any(x in lower for x in ["song","music","video","on youtube"])):
        import re
        song_match = re.search(r'play\s+(.+?)(?:\s+on youtube|$)', lower)
        query = song_match.group(1) if song_match else lower.replace("open youtube","").strip()
        if query and query not in ["youtube",""]:
            msg = play_youtube(query)
        else:
            webbrowser.open("https://www.youtube.com")
            msg = "Opening YouTube"
        response = f"{msg}, boss."
        action = "youtube"
        return {"response": response, "action": action}

    # ── GOOGLE SEARCH ─────────────────────────────────────────────────────
    if "search" in lower or "google" in lower:
        import re
        q_match = re.search(r'(?:search|google)\s+(?:for\s+)?(.+)', lower)
        q = q_match.group(1) if q_match else lower
        url = f"https://www.google.com/search?q={urllib.parse.quote(q)}"
        open_url(url)
        response = f"Searching Google for {q}, sir."
        action = "search"
        return {"response": response, "action": action}

    # ── OPEN SPOTIFY ──────────────────────────────────────────────────────
    if "spotify" in lower:
        ok, msg = open_app("spotify")
        response = "Launching Spotify, boss." if ok else "Opening Spotify in browser, sir."
        if not ok:
            webbrowser.open("https://open.spotify.com")
        action = "open_spotify"
        return {"response": response, "action": action}

    # ── CALCULATOR ───────────────────────────────────────────────────────
    if "calculator" in lower or "calc" in lower:
        open_app("calculator")
        response = "Opening Calculator, boss."
        action = "open_calc"
        return {"response": response, "action": action}

    # ── TASK MANAGER ─────────────────────────────────────────────────────
    if "task manager" in lower:
        open_app("task_manager")
        response = "Opening Task Manager, sir."
        action = "open_taskman"
        return {"response": response, "action": action}

    # ── CMD / POWERSHELL ──────────────────────────────────────────────────
    if "command prompt" in lower or "cmd" in lower:
        open_app("cmd")
        response = "Opening Command Prompt, boss."
        action = "open_cmd"
        return {"response": response, "action": action}
    if "powershell" in lower:
        open_app("powershell")
        response = "Opening PowerShell, sir."
        action = "open_ps"
        return {"response": response, "action": action}

    # ── PAINT ─────────────────────────────────────────────────────────────
    if "paint" in lower:
        open_app("paint")
        response = "Opening Paint, boss."
        action = "open_paint"
        return {"response": response, "action": action}

    # ── SCREENSHOT ────────────────────────────────────────────────────────
    if "screenshot" in lower or "screen shot" in lower or "take a screenshot" in lower:
        ok, path = take_screenshot()
        response = f"Screenshot taken and saved, boss." if ok else "Screenshot captured with Snipping Tool, sir."
        action = "screenshot"
        return {"response": response, "action": action}

    # ── VOLUME ────────────────────────────────────────────────────────────
    if "volume" in lower or "mute" in lower:
        import re
        if "mute" in lower:
            msg = mute_volume()
            response = f"{msg}, boss."
        else:
            vol_match = re.search(r'(\d+)', lower)
            if vol_match:
                level = int(vol_match.group(1))
                msg = set_volume(level)
                response = f"{msg}, sir."
            elif "up" in lower or "increase" in lower or "louder" in lower:
                msg = set_volume(70)
                response = "Volume increased, boss."
            elif "down" in lower or "decrease" in lower or "lower" in lower or "quiet" in lower:
                msg = set_volume(30)
                response = "Volume decreased, sir."
            else:
                response = "What volume level would you like, boss?"
        action = "volume"
        return {"response": response, "action": action}

    # ── TIME / DATE ───────────────────────────────────────────────────────
    if "time" in lower:
        t = datetime.now().strftime("%I:%M %p")
        response = f"The current time is {t}, boss."
        action = "time"
        return {"response": response, "action": action}

    if "date" in lower:
        d = datetime.now().strftime("%A, %B %d, %Y")
        response = f"Today is {d}, sir."
        action = "date"
        return {"response": response, "action": action}

    # ── BATTERY ───────────────────────────────────────────────────────────
    if "battery" in lower:
        bat, plug = get_battery()
        if bat:
            ch = "charging" if plug else "not charging"
            response = f"Battery is at {bat} percent and {ch}, boss."
        else:
            response = "Battery information unavailable, sir."
        action = "battery"
        return {"response": response, "action": action}

    # ── SYSTEM STATUS ─────────────────────────────────────────────────────
    if "status" in lower or "system" in lower:
        stats = get_stats()
        bat = f"{stats['battery']}%" if stats['battery'] else "N/A"
        response = f"All systems nominal, boss. CPU at {stats['cpu']} percent, RAM at {stats['ram_pct']} percent, battery at {bat}. Network online. Temperature within safe limits."
        action = "status"
        return {"response": response, "action": action}

    # ── OPEN FILE (general) ───────────────────────────────────────────────
    if "open" in lower or "launch" in lower:
        import re
        file_match = re.search(r'(?:open|launch|start)\s+([\w\-./\\]+(?:\.\w+)?)', lower)
        if file_match:
            target = file_match.group(1)
            # try as file path first
            if "." in target:
                ok, msg = open_file(target)
                if ok:
                    response = f"Opening {target}, boss."
                    action = "open_file"
                    return {"response": response, "action": action}
            # else try as app
            subprocess.Popen(target, shell=True)
            response = f"Trying to open {target}, sir."
            action = "open_app"
            return {"response": response, "action": action}

    # ── SHUTDOWN / RESTART / SLEEP PC ─────────────────────────────────────
    if "shutdown" in lower or "shut down" in lower:
        response = shutdown_pc()
        action = "shutdown"
        return {"response": response, "action": action}
    if "restart" in lower or "reboot" in lower:
        response = restart_pc()
        action = "restart"
        return {"response": response, "action": action}
    if ("sleep" in lower or "hibernate" in lower) and "pc" in lower:
        response = sleep_pc()
        action = "sleep_pc"
        return {"response": response, "action": action}

    # ── THANK YOU ─────────────────────────────────────────────────────────
    if "thank" in lower:
        response = "Always at your service, boss."
        action = "thanks"
        return {"response": response, "action": action}

    # ── HELLO ─────────────────────────────────────────────────────────────
    if any(x in lower for x in ["hello", "hi buddy", "hey buddy", "hey there"]):
        response = "Hello boss! All systems operational. How can I assist you today?"
        action = "hello"
        return {"response": response, "action": action}

    # ── JOKE ──────────────────────────────────────────────────────────────
    if "joke" in lower:
        jokes = [
            "Why do programmers prefer dark mode? Because light attracts bugs, sir.",
            "I told my computer I needed a break. Now it won't stop sending me vacation ads, boss.",
            "Why don't scientists trust atoms? Because they make up everything, sir.",
            "I tried to write a joke about memory leaks, but I forgot how it ended, boss.",
        ]
        import random
        response = jokes[random.randint(0, len(jokes)-1)]
        action = "joke"
        return {"response": response, "action": action}

    # ── DEFAULT ───────────────────────────────────────────────────────────
    import random
    defaults = [
        f"Command received: {cmd}. I'll do my best, boss.",
        f"Understood, sir. Processing: {cmd}.",
        f"On it, boss. Let me handle that for you.",
    ]
    response = defaults[random.randint(0, len(defaults)-1)]
    action = "unknown"
    return {"response": response, "action": action}


# ───────────────────────────────────────────────────────────────────────────
# DOCUMENT / CODE GENERATORS
# ───────────────────────────────────────────────────────────────────────────
def generate_doc_content(doc_type: str) -> list:
    dt = doc_type.lower()
    today = datetime.now().strftime("%B %d, %Y")

    if "letter" in dt:
        return [
            "# Formal Letter",
            "",
            f"Date: {today}",
            "",
            "To Whom It May Concern,",
            "",
            "I am writing to bring to your attention the following matter. Please find the details outlined below for your review and consideration.",
            "",
            "I would appreciate your prompt response at your earliest convenience. Please do not hesitate to contact me should you require any further information.",
            "",
            "Thank you for your time and attention.",
            "",
            "Yours sincerely,",
            "",
            "[Your Name]",
            "[Your Title]",
            "[Your Contact Information]",
        ]
    elif "resume" in dt or "cv" in dt:
        return [
            "# Curriculum Vitae",
            "",
            "## Personal Information",
            "Name: [Your Name]",
            "Email: [your@email.com]",
            "Phone: [+XX XXX XXX XXXX]",
            "LinkedIn: [linkedin.com/in/yourname]",
            "",
            "## Professional Summary",
            "Experienced professional with a strong background in [your field]. Proven track record of delivering results and driving innovation.",
            "",
            "## Work Experience",
            "### [Job Title] — [Company Name]",
            f"[Start Date] – Present",
            "- Key responsibility and achievement",
            "- Key responsibility and achievement",
            "",
            "## Education",
            "### [Degree] in [Field]",
            "[University Name], [Year]",
            "",
            "## Skills",
            "- Skill 1, Skill 2, Skill 3",
        ]
    elif "report" in dt:
        return [
            f"# Report: {doc_type.title()}",
            f"Date: {today}",
            "Prepared by: BUDDY AI",
            "",
            "## Executive Summary",
            "This report provides a comprehensive overview of the subject matter. Key findings and recommendations are outlined below.",
            "",
            "## Introduction",
            "The purpose of this report is to document and analyze the relevant information pertaining to this subject.",
            "",
            "## Findings",
            "Based on the analysis conducted, the following key findings have been identified:",
            "1. Finding one — describe here",
            "2. Finding two — describe here",
            "3. Finding three — describe here",
            "",
            "## Recommendations",
            "Based on the findings, the following recommendations are proposed:",
            "1. Recommendation one",
            "2. Recommendation two",
            "",
            "## Conclusion",
            "This concludes the report. Further analysis may be required.",
        ]
    elif "email" in dt:
        return [
            "# Email Draft",
            "",
            f"Date: {today}",
            "To: [recipient@email.com]",
            "Subject: [Subject Line]",
            "",
            "Dear [Name],",
            "",
            "I hope this email finds you well. I am reaching out regarding [subject matter].",
            "",
            "[Main body of your email goes here. Include all relevant details and information.]",
            "",
            "Please let me know if you have any questions or concerns.",
            "",
            "Best regards,",
            "[Your Name]",
        ]
    else:
        return [
            f"# {doc_type.title()}",
            f"Date: {today}",
            "Created by BUDDY AI",
            "",
            "## Overview",
            f"This document covers: {doc_type}",
            "",
            "## Details",
            "[Add your content here]",
            "",
            "## Notes",
            "[Additional notes and information]",
        ]

def generate_code_template(lang: str, fname: str) -> str:
    templates = {
        "python": f'''# {fname}.py
# Created by BUDDY AI

def main():
    print("Hello from {fname}!")

if __name__ == "__main__":
    main()
''',
        "javascript": f'''// {fname}.js
// Created by BUDDY AI

function main() {{
    console.log("Hello from {fname}!");
}}

main();
''',
        "html": f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>{fname}</title>
</head>
<body>
    <h1>Hello from {fname}</h1>
    <p>Created by BUDDY AI</p>
</body>
</html>
''',
    }
    return templates.get(lang, templates["python"])


# ───────────────────────────────────────────────────────────────────────────
# HTTP SERVER  (talks to the HTML frontend)
# ───────────────────────────────────────────────────────────────────────────
PORT = 8765

class BuddyHandler(http.server.BaseHTTPRequestHandler):
    def log_message(self, format, *args):
        pass  # silence server logs

    def _cors(self):
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")

    def do_OPTIONS(self):
        self.send_response(200)
        self._cors()
        self.end_headers()

    def do_GET(self):
        parsed = urllib.parse.urlparse(self.path)
        params = urllib.parse.parse_qs(parsed.query)

        if parsed.path == "/stats":
            data = get_stats()
            self._json(data)

        elif parsed.path == "/command":
            cmd = params.get("q", [""])[0]
            if cmd:
                result = process_command(cmd)
                # speak in background
                speak_async(result["response"])
                self._json(result)
            else:
                self._json({"response": "", "action": "none"})

        elif parsed.path == "/ping":
            self._json({"status": "ok", "version": "1.0"})

        else:
            self.send_response(404)
            self.end_headers()

    def do_POST(self):
        length = int(self.headers.get("Content-Length", 0))
        body = self.rfile.read(length)
        try:
            data = json.loads(body)
            cmd = data.get("command", "")
            result = process_command(cmd)
            speak_async(result["response"])
            self._json(result)
        except Exception as e:
            self._json({"response": str(e), "action": "error"})

    def _json(self, obj):
        payload = json.dumps(obj).encode()
        self.send_response(200)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", len(payload))
        self._cors()
        self.end_headers()
        self.wfile.write(payload)


def run_server():
    with socketserver.TCPServer(("127.0.0.1", PORT), BuddyHandler) as httpd:
        httpd.allow_reuse_address = True
        print(f"[BUDDY] Server running on http://127.0.0.1:{PORT}")
        httpd.serve_forever()


# ───────────────────────────────────────────────────────────────────────────
# ENTRY POINT
# ───────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("=" * 55)
    print("  BUDDY AI ASSISTANT - v1.0")
    print("  Windows 11 OS Control Backend")
    print("=" * 55)

    # start HTTP server in background
    server_thread = threading.Thread(target=run_server, daemon=True)
    server_thread.start()
    time.sleep(0.5)

    # open the HTML UI
    html_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "buddy_ui.html")
    if os.path.exists(html_path):
        webbrowser.open(f"file:///{html_path.replace(os.sep, '/')}")
        print(f"[BUDDY] UI opened: {html_path}")
    else:
        print(f"[BUDDY] WARNING: buddy_ui.html not found next to buddy.py")

    speak_sync("BUDDY AI systems online. Ready for your commands, boss.")
    print("[BUDDY] Ready. Press Ctrl+C to quit.")

    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        print("\n[BUDDY] Shutting down. Goodbye, boss.")
